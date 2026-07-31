#!/usr/bin/env python3
"""Build the deterministic public fixtures and metadata for the v0.2 release.

The build has two phases so the checked-in JavaScript office builder can run in
an @oai/artifact-tool workspace between them:

    python3 tools/artifacts/build_release.py --phase prepare
    node tools/artifacts/build_office_assets.mjs <repo-root> <qa-root>
    python3 tools/artifacts/build_release.py --phase finalize
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import io
import json
import re
import shutil
import zipfile
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[2]
DATA_ROOT = ROOT / "data" / "v0.2"
ASSET_ROOT = DATA_ROOT / "assets"
CALIBRATION_ROOT = DATA_ROOT / "public-calibration"
SCENARIO_ROOT = DATA_ROOT / "scenarios"
REPORT_PATH = ROOT / "reports" / "v0.2" / "release-validation.json"
BUILD_DATE = "2026-07-31"

DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
PPTX_MIME = (
    "application/vnd.openxmlformats-officedocument.presentationml.presentation"
)
XLSX_MIME = (
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)

MODULES = [
    "attached_file_grounding",
    "document_work",
    "presentation_work",
    "spreadsheet_work",
    "cross_artifact_workflow",
    "artifact_quality_control_delivery",
]
WORKFLOW_JOBS = [
    "inspect",
    "extract",
    "synthesize",
    "create",
    "revise",
    "repurpose",
    "validate",
    "package",
]


def _write_text(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value.rstrip() + "\n", encoding="utf-8")


def _write_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(value, ensure_ascii=False, indent=2, sort_keys=False) + "\n",
        encoding="utf-8",
    )


def _write_csv(path: Path, header: list[str], rows: list[list[Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle, lineterminator="\n")
        writer.writerow(header)
        writer.writerows(rows)


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _relative(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()


def _set_run_font(
    run: Any,
    *,
    name: str = "Calibri",
    size: float = 11,
    color: str = "182230",
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def _set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", 100), ("start", 120), ("bottom", 100), ("end", 120)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    table_width = table_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    table_indent = table_pr.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(widths[index]))
            tc_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_page_field(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    _set_run_font(run, size=9, color="5D6B79")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def _new_business_document(title: str, subtitle: str, kicker: str) -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2F6B8A", 14, 7),
        ("Heading 2", 13, "2F6B8A", 10, 5),
        ("Heading 3", 12, "17324D", 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("RWPK v0.2 | Synthetic calibration artifact")
    _set_run_font(run, size=9, color="5D6B79")
    _add_page_field(section.footer.paragraphs[0])

    kicker_paragraph = document.add_paragraph()
    kicker_paragraph.paragraph_format.space_after = Pt(4)
    run = kicker_paragraph.add_run(kicker.upper())
    _set_run_font(run, size=9.5, color="2F6B8A", bold=True)

    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(4)
    run = title_paragraph.add_run(title)
    _set_run_font(run, size=23, color="17324D", bold=True)

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.paragraph_format.space_after = Pt(14)
    run = subtitle_paragraph.add_run(subtitle)
    _set_run_font(run, size=12.5, color="5D6B79")

    metadata = document.add_paragraph()
    metadata.paragraph_format.space_after = Pt(12)
    label = metadata.add_run("Prepared: ")
    _set_run_font(label, size=9.5, color="5D6B79", bold=True)
    value = metadata.add_run(BUILD_DATE)
    _set_run_font(value, size=9.5, color="5D6B79")
    return document


def _add_bullets(document: Document, bullets: list[str]) -> None:
    for item in bullets:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.line_spacing = 1.1
        run = paragraph.add_run(item)
        _set_run_font(run)


def _add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int] | None = None,
) -> None:
    if widths is None:
        base = 9360 // len(headers)
        widths = [base] * len(headers)
        widths[-1] += 9360 - sum(widths)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, "2F6B8A")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
        )
        run = paragraph.add_run(header)
        _set_run_font(run, size=9.5, color="FFFFFF", bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            run = paragraph.add_run(str(value))
            _set_run_font(run, size=9.5)
    _set_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _save_document(document: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document.core_properties.author = "RealWorld Prompt Kit"
    document.core_properties.last_modified_by = "RealWorld Prompt Kit"
    fixed = datetime(2026, 7, 31, 0, 0, 0, tzinfo=UTC)
    document.core_properties.created = fixed
    document.core_properties.modified = fixed
    document.save(path)


def _build_document(
    path: Path,
    *,
    title: str,
    subtitle: str,
    kicker: str,
    sections: list[dict[str, Any]],
) -> None:
    document = _new_business_document(title, subtitle, kicker)
    for section in sections:
        document.add_heading(section["heading"], level=1)
        for paragraph_text in section.get("paragraphs", []):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(paragraph_text)
            _set_run_font(run)
        if section.get("bullets"):
            _add_bullets(document, section["bullets"])
        if section.get("table"):
            table = section["table"]
            _add_table(
                document,
                table["headers"],
                table["rows"],
                table.get("widths"),
            )
    _save_document(document, path)


def _build_vendor_pdf(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    story: list[Any] = [
        Paragraph("Facilities support vendor packet", styles["Title"]),
        Spacer(1, 0.2 * inch),
        Paragraph(
            "Synthetic quotes for a 12-month office support contract.",
            styles["BodyText"],
        ),
        Spacer(1, 0.2 * inch),
    ]
    table = Table(
        [
            ["Vendor", "Annual quote", "Lead time", "Service window"],
            ["Northline", "$52,000", "4 weeks", "Weekdays 08:00-18:00"],
            ["Cobalt", "$48,500", "6 weeks", "Weekdays 07:00-19:00"],
            ["Harbor", "$55,000", "3 weeks", "24/5"],
        ],
        colWidths=[1.2 * inch, 1.3 * inch, 1.1 * inch, 2.6 * inch],
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F6B8A")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D7E1E7")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EDF4F7")]),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.extend(
        [
            table,
            Spacer(1, 0.3 * inch),
            Paragraph(
                "<b>Revision note:</b> Cobalt submitted a revised annual quote of "
                "$51,000 after the comparison table was prepared. Treat $48,500 "
                "and $51,000 as a source conflict; do not silently select one.",
                styles["BodyText"],
            ),
            Spacer(1, 0.2 * inch),
            Paragraph(
                "Required decision date: 2026-08-07. Contract authority remains "
                "with the procurement lead.",
                styles["BodyText"],
            ),
        ]
    )
    document = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=0.8 * inch,
        leftMargin=0.8 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        invariant=1,
        title="Facilities support vendor packet",
        author="RealWorld Prompt Kit",
    )
    document.build(story)


def _load_scan_font(size: int) -> ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def _build_scanned_expense_pdf(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1800, 1200), "#F5F1E8")
    draw = ImageDraw.Draw(image)
    title_font = _load_scan_font(52)
    header_font = _load_scan_font(28)
    body_font = _load_scan_font(27)
    draw.text((90, 64), "APRIL EXPENSE REGISTER - SCANNED COPY", fill="#1C2733", font=title_font)
    headers = ["Date", "Department", "Category", "Vendor", "Amount", "Receipt ID"]
    rows = [
        ["2026-04-03", "Operations", "Shipping", "Cobalt Freight", "$1,280", "R-1041"],
        ["2026-04-06", "People", "Training", "Northline Learning", "$860", "R-1042"],
        ["2026-04-09", "Sales", "Travel", "Harbor Rail", "$420", "R-1043"],
        ["2026-04-13", "Operations", "Supplies", "Beacon Office", "$315", "R-1044"],
        ["2026-04-18", "Marketing", "Events", "Civic Hall", "$2,100", "R-1045"],
        ["2026-04-24", "People", "Recruiting", "Evergreen Jobs", "$740", "R-1046"],
    ]
    widths = [210, 260, 220, 350, 180, 190]
    x_positions = [70]
    for width in widths:
        x_positions.append(x_positions[-1] + width)
    top = 180
    row_height = 105
    for row_index in range(len(rows) + 2):
        y = top + row_index * row_height
        draw.line((70, y, x_positions[-1], y), fill="#5C6770", width=3)
    for x in x_positions:
        draw.line((x, top, x, top + (len(rows) + 1) * row_height), fill="#5C6770", width=3)
    for index, header in enumerate(headers):
        draw.text((x_positions[index] + 12, top + 28), header, fill="#1C2733", font=header_font)
    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row):
            draw.text(
                (x_positions[col_index] + 12, top + row_index * row_height + 30),
                value,
                fill="#303840",
                font=body_font,
            )
    draw.text(
        (82, 1060),
        "Synthetic source | verify all six rows and the amount total",
        fill="#5C6770",
        font=body_font,
    )
    buffer = io.BytesIO()
    image.save(buffer, format="PNG", optimize=False)
    buffer.seek(0)

    pdf = canvas.Canvas(
        str(path),
        pagesize=letter,
        invariant=1,
        pageCompression=1,
    )
    pdf.setTitle("Scanned expense table")
    pdf.setAuthor("RealWorld Prompt Kit")
    pdf.drawInlineImage(
        Image.open(buffer),
        0.45 * inch,
        1.1 * inch,
        width=7.6 * inch,
        height=5.07 * inch,
    )
    pdf.showPage()
    pdf.save()


def _normalize_ooxml(path: Path) -> None:
    if path.suffix.lower() not in {".docx", ".pptx", ".xlsx"}:
        return
    with zipfile.ZipFile(path, "r") as source:
        entries = [(info.filename, source.read(info.filename)) for info in source.infolist()]
    relationship_ids: dict[str, str] = {}
    creation_guids: dict[str, str] = {}
    creation_numbers: dict[str, str] = {}

    def canonical_relationship(match: re.Match[str]) -> str:
        value = match.group(0)
        if value not in relationship_ids:
            relationship_ids[value] = f"R{len(relationship_ids) + 1:016x}"
        return relationship_ids[value]

    def canonical_guid(match: re.Match[str]) -> str:
        value = match.group(0)
        if value not in creation_guids:
            token = f"{len(creation_guids) + 1:032X}"
            creation_guids[value] = (
                f"{{{token[:8]}-{token[8:12]}-{token[12:16]}-"
                f"{token[16:20]}-{token[20:]}}}"
            )
        return creation_guids[value]

    def canonical_creation_number(match: re.Match[str]) -> str:
        value = match.group(0)
        if value not in creation_numbers:
            creation_numbers[value] = str(len(creation_numbers) + 1)
        return creation_numbers[value]

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as target:
        for filename, payload in sorted(entries):
            if filename.endswith((".xml", ".rels")):
                text = payload.decode("utf-8-sig")
                if filename == "docProps/core.xml":
                    text = re.sub(
                        r"(<dcterms:(?:created|modified)[^>]*>)[^<]*(</dcterms:(?:created|modified)>)",
                        rf"\g<1>{BUILD_DATE}T00:00:00Z\g<2>",
                        text,
                    )
                text = re.sub(
                    r"R[0-9A-Fa-f]{16}",
                    canonical_relationship,
                    text,
                )
                text = re.sub(
                    r'(?<=<a16:creationId id=")\{[0-9A-Fa-f-]{36}\}(?=")',
                    canonical_guid,
                    text,
                )
                text = re.sub(
                    r'(?<=<p14:creationId val=")[0-9]+(?=")',
                    canonical_creation_number,
                    text,
                )
                payload = text.encode("utf-8")
            info = zipfile.ZipInfo(filename, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            target.writestr(info, payload)
    path.write_bytes(buffer.getvalue())


def _prepare_text_and_csv_assets() -> None:
    _write_text(
        ASSET_ROOT
        / "rwpk.communication_collaboration_negotiation.meeting_minutes.0003"
        / "meeting-notes.txt",
        """
Weekly program sync - 2026-07-28

Decision: keep the pilot launch on 2026-08-17.
Decision: the data migration dry run must finish before training starts.
Action: Operations role owns the dry run by 2026-08-05.
Action: Enablement role drafts the quick guide by 2026-08-07.
Open question: Support escalation coverage for launch week is not assigned.
Risk: two pilot locations have not confirmed their local champions.
""",
    )
    _write_text(
        ASSET_ROOT
        / "rwpk.transformation_rewriting.onboarding_sop_revision.0004"
        / "reviewer-feedback.txt",
        """
Apply only these changes:
1. Replace the 48-hour access target with 24 hours.
2. Assign access approval to the Hiring Manager role, not People Operations.
3. Add a final verification step owned by IT Service Desk.
4. Keep the existing scope and do not add software procurement.
""",
    )
    _write_csv(
        ASSET_ROOT
        / "rwpk.summarization_synthesis.executive_update_deck.0005"
        / "workstream-metrics.csv",
        ["workstream", "completion_pct", "status", "decision_needed"],
        [
            ["Platform", 94, "On plan", "No"],
            ["Data", 91, "On plan", "No"],
            ["Enablement", 68, "At risk", "Capacity"],
            ["Support", 87, "On plan", "Owner"],
        ],
    )
    _write_text(
        ASSET_ROOT
        / "rwpk.transformation_rewriting.deck_feedback_revision.0006"
        / "reviewer-comments.txt",
        """
Revise the existing deck without changing its source figures.
- Separate verified readiness from owner-reported confidence.
- Make handoff ownership the central risk.
- Keep the deck to three slides.
- Preserve the Q2 2026 period and do not introduce a launch date.
""",
    )
    _write_csv(
        ASSET_ROOT
        / "rwpk.quantitative_formal_analysis.project_status_workbook.0007"
        / "project-status.csv",
        ["project", "owner_role", "status", "budget", "budget_used", "target_date"],
        [
            ["Atlas migration", "Operations", "On Track", 42000, 27400, "2026-08-15"],
            ["Beacon onboarding", "Customer Success", "At Risk", 28000, 24100, "2026-07-31"],
            ["Cobalt launch", "Marketing", "On Track", 35000, 18200, "2026-09-10"],
            ["Delta controls", "Finance", "Blocked", 18000, 9600, "2026-08-01"],
            ["Evergreen renewal", "Sales", "At Risk", 22000, 19600, "2026-08-20"],
            ["Harbor training", "People", "On Track", 12000, 5400, "2026-09-01"],
        ],
    )
    _write_csv(
        ASSET_ROOT
        / "rwpk.operations_monitoring_improvement.monthly_ops_package.0009"
        / "monthly-ops.csv",
        ["month", "orders", "on_time_rate", "backlog", "support_cases"],
        [
            ["Jan", 1180, 0.91, 84, 126],
            ["Feb", 1240, 0.92, 79, 118],
            ["Mar", 1310, 0.90, 96, 134],
            ["Apr", 1390, 0.93, 71, 112],
            ["May", 1460, 0.94, 63, 105],
            ["Jun", 1525, 0.95, 54, 99],
        ],
    )
    _write_text(
        ASSET_ROOT
        / "rwpk.operations_monitoring_improvement.monthly_ops_package.0009"
        / "management-brief.txt",
        """
Prepare a management package for the January-June 2026 operating review.
Keep the source rows in the workbook. Highlight the relationship between
orders, on-time delivery, and backlog. The memo and deck must use the same June
figures and identify the 93% service floor as a monitoring threshold, not a
historical target.
""",
    )
    _write_csv(
        ASSET_ROOT
        / "rwpk.evaluation_review_audit.quarterly_report_audit.0011"
        / "quarterly-source.csv",
        ["region", "q1_orders", "q2_orders", "q2_on_time_rate"],
        [
            ["North", 820, 860, 0.95],
            ["South", 760, 735, 0.91],
            ["East", 690, 745, 0.94],
            ["West", 710, 772, 0.93],
        ],
    )
    _write_csv(
        ASSET_ROOT
        / "rwpk.one_off_tool_execution.delivery_package.0012"
        / "delivery-data.csv",
        ["deliverable", "owner_role", "status", "due_date"],
        [
            ["Executive recap", "Program lead", "Complete", "2026-07-31"],
            ["Action tracker", "Operations", "Complete", "2026-07-31"],
            ["Handoff note", "Program lead", "Complete", "2026-07-31"],
            ["Archive copy", "Operations", "Ready", "2026-08-03"],
        ],
    )


def _prepare_documents() -> None:
    _build_document(
        ASSET_ROOT
        / "rwpk.communication_collaboration_negotiation.meeting_minutes.0003"
        / "meeting-agenda.docx",
        title="Weekly program sync",
        subtitle="Agenda and decision frame",
        kicker="Source agenda",
        sections=[
            {
                "heading": "Objectives",
                "bullets": [
                    "Confirm the pilot launch date.",
                    "Sequence migration and training dependencies.",
                    "Assign unresolved launch-week coverage.",
                ],
            },
            {
                "heading": "Agenda",
                "table": {
                    "headers": ["Topic", "Outcome"],
                    "rows": [
                        ["Schedule", "Decision"],
                        ["Migration", "Owner and due date"],
                        ["Enablement", "Owner and due date"],
                        ["Support", "Open question"],
                    ],
                    "widths": [2600, 6760],
                },
            },
        ],
    )
    _build_document(
        ASSET_ROOT
        / "rwpk.transformation_rewriting.onboarding_sop_revision.0004"
        / "onboarding-sop-draft.docx",
        title="New starter access SOP",
        subtitle="Draft requiring a scoped policy revision",
        kicker="Existing artifact",
        sections=[
            {
                "heading": "Purpose",
                "paragraphs": [
                    "Provide system access to new starters within 48 hours while preserving approval evidence."
                ],
            },
            {
                "heading": "Procedure",
                "table": {
                    "headers": ["Step", "Owner", "Target"],
                    "rows": [
                        ["Submit access request", "Hiring Manager", "Before start date"],
                        ["Approve access", "People Operations", "Within 48 hours"],
                        ["Provision accounts", "IT Service Desk", "After approval"],
                    ],
                    "widths": [3900, 2700, 2760],
                },
            },
            {
                "heading": "Scope",
                "paragraphs": [
                    "This SOP covers standard identity, collaboration, and reporting access. Software procurement is excluded."
                ],
            },
        ],
    )
    _build_document(
        ASSET_ROOT
        / "rwpk.summarization_synthesis.executive_update_deck.0005"
        / "portfolio-report.docx",
        title="Portfolio delivery report",
        subtitle="Source narrative for the June executive update",
        kicker="Authoritative source",
        sections=[
            {
                "heading": "Executive context",
                "paragraphs": [
                    "Three of four workstreams remain on plan. Enablement is constrained by facilitator capacity, while Support still needs a named post-launch escalation owner."
                ],
            },
            {
                "heading": "Decision required",
                "bullets": [
                    "Approve one additional facilitator through 2026-07-18.",
                    "Keep the Platform change freeze until migration validation closes.",
                    "Name the Support escalation owner before launch readiness review.",
                ],
            },
        ],
    )
    _build_document(
        ASSET_ROOT
        / "rwpk.evaluation_review_audit.campaign_readout_package.0010"
        / "brand-guide.docx",
        title="Campaign readout style guide",
        subtitle="Synthetic communication and visual direction",
        kicker="Style reference",
        sections=[
            {
                "heading": "Voice",
                "bullets": [
                    "Lead with a concrete finding, then show evidence.",
                    "Use direct language and avoid promotional superlatives.",
                    "Distinguish scale from efficiency.",
                ],
            },
            {
                "heading": "Visual direction",
                "table": {
                    "headers": ["Role", "Color", "Use"],
                    "rows": [
                        ["Primary", "#17324D", "Titles and anchors"],
                        ["Secondary", "#2F6B8A", "Charts and section cues"],
                        ["Accent", "#57C7D4", "Selective emphasis"],
                    ],
                    "widths": [2200, 2200, 4960],
                },
            },
        ],
    )
    _build_document(
        ASSET_ROOT
        / "rwpk.evaluation_review_audit.quarterly_report_audit.0011"
        / "quarterly-report-draft.docx",
        title="Quarterly regional operations report",
        subtitle="Draft for source and quality audit",
        kicker="Existing artifact",
        sections=[
            {
                "heading": "Draft summary",
                "paragraphs": [
                    "All regions improved order volume in Q2, and every region exceeded a 94% on-time rate."
                ],
            },
            {
                "heading": "Draft recommendations",
                "bullets": [
                    "Reuse the North region playbook in every region.",
                    "Increase South region capacity immediately.",
                ],
            },
        ],
    )
    _build_document(
        ASSET_ROOT
        / "rwpk.one_off_tool_execution.delivery_package.0012"
        / "delivery-brief.docx",
        title="Quarterly handoff package brief",
        subtitle="Required files, names, and authority boundary",
        kicker="Delivery contract source",
        sections=[
            {
                "heading": "Required package",
                "table": {
                    "headers": ["File", "Purpose"],
                    "rows": [
                        ["delivery-recap-deck.pptx", "Executive recap"],
                        ["delivery-tracker.xlsx", "Editable owner tracker"],
                        ["handoff-note.docx", "Handoff summary"],
                        ["delivery-manifest.json", "Machine-readable file list"],
                    ],
                    "widths": [3800, 5560],
                },
            },
            {
                "heading": "Boundary",
                "paragraphs": [
                    "Write only to the local output package. Preserve all supplied inputs and do not send, upload, or archive anything."
                ],
            },
        ],
    )

    _build_document(
        CALIBRATION_ROOT
        / "rwpk.summarization_synthesis.vendor_packet_brief.0001"
        / "vendor-comparison-brief.docx",
        title="Facilities vendor comparison",
        subtitle="Decision brief with an explicit source conflict",
        kicker="Calibration reference",
        sections=[
            {
                "heading": "Decision summary",
                "paragraphs": [
                    "Cobalt cannot be ranked on price until procurement resolves the conflict between the $48,500 table value and the revised $51,000 quote."
                ],
            },
            {
                "heading": "Comparable facts",
                "table": {
                    "headers": ["Vendor", "Price evidence", "Lead time", "Service"],
                    "rows": [
                        ["Northline", "$52,000", "4 weeks", "Weekdays 08:00-18:00"],
                        ["Cobalt", "$48,500 / $51,000 conflict", "6 weeks", "Weekdays 07:00-19:00"],
                        ["Harbor", "$55,000", "3 weeks", "24/5"],
                    ],
                    "widths": [1900, 3000, 1700, 2760],
                },
            },
            {
                "heading": "Recommended next step",
                "bullets": [
                    "Hold the price-based recommendation.",
                    "Ask procurement to confirm the controlling Cobalt quote before 2026-08-07.",
                    "Retain lead time and service-window differences for the final decision.",
                ],
            },
        ],
    )
    _build_document(
        CALIBRATION_ROOT
        / "rwpk.communication_collaboration_negotiation.meeting_minutes.0003"
        / "program-sync-minutes.docx",
        title="Weekly program sync minutes",
        subtitle="Decisions, owners, dates, and unresolved questions",
        kicker="Calibration reference",
        sections=[
            {
                "heading": "Decisions",
                "bullets": [
                    "Keep the pilot launch on 2026-08-17.",
                    "Complete the data migration dry run before training starts.",
                ],
            },
            {
                "heading": "Actions",
                "table": {
                    "headers": ["Action", "Owner role", "Due date", "Status"],
                    "rows": [
                        ["Complete migration dry run", "Operations", "2026-08-05", "Open"],
                        ["Draft quick guide", "Enablement", "2026-08-07", "Open"],
                    ],
                    "widths": [3900, 1900, 1700, 1860],
                },
            },
            {
                "heading": "Unresolved questions",
                "bullets": [
                    "Who owns Support escalation coverage for launch week?",
                    "Which two pilot locations still need local champion confirmation?",
                ],
            },
        ],
    )
    _build_document(
        CALIBRATION_ROOT
        / "rwpk.transformation_rewriting.onboarding_sop_revision.0004"
        / "onboarding-sop-revised.docx",
        title="New starter access SOP",
        subtitle="Scoped revision with a 24-hour target and verification step",
        kicker="Calibration reference",
        sections=[
            {
                "heading": "Purpose",
                "paragraphs": [
                    "Provide standard system access to new starters within 24 hours while preserving approval and verification evidence."
                ],
            },
            {
                "heading": "Procedure",
                "table": {
                    "headers": ["Step", "Owner", "Target"],
                    "rows": [
                        ["Submit access request", "Hiring Manager", "Before start date"],
                        ["Approve access", "Hiring Manager", "Within 24 hours"],
                        ["Provision accounts", "IT Service Desk", "After approval"],
                        ["Verify access and record result", "IT Service Desk", "Before closure"],
                    ],
                    "widths": [3900, 2700, 2760],
                },
            },
            {
                "heading": "Scope",
                "paragraphs": [
                    "This SOP covers standard identity, collaboration, and reporting access. Software procurement remains excluded."
                ],
            },
        ],
    )
    _build_document(
        CALIBRATION_ROOT
        / "rwpk.operations_monitoring_improvement.monthly_ops_package.0009"
        / "monthly-ops-memo.docx",
        title="Monthly operations memo",
        subtitle="June closed with higher throughput and a smaller backlog",
        kicker="Calibration reference",
        sections=[
            {
                "heading": "Executive readout",
                "paragraphs": [
                    "June orders reached 1,525, on-time delivery reached 95%, and backlog fell to 54. This is the strongest combined service result in the six-month source."
                ],
            },
            {
                "heading": "Monitoring thresholds",
                "bullets": [
                    "Escalate if on-time delivery falls below the 93% management floor.",
                    "Review workload if backlog rises above 70 items.",
                    "Keep the source workbook as the editable calculation record.",
                ],
            },
        ],
    )
    _build_document(
        CALIBRATION_ROOT
        / "rwpk.evaluation_review_audit.campaign_readout_package.0010"
        / "campaign-readout-brief.docx",
        title="Campaign performance readout",
        subtitle="Scale and efficiency separated for the next-cycle decision",
        kicker="Calibration reference",
        sections=[
            {
                "heading": "Finding",
                "paragraphs": [
                    "Partner produced the most efficient qualified demand at approximately $130 per qualified lead. Search generated the largest pipeline at $310,000."
                ],
            },
            {
                "heading": "Channel comparison",
                "table": {
                    "headers": ["Channel", "Spend", "Qualified", "Pipeline", "Decision role"],
                    "rows": [
                        ["Search", "$24,000", "198", "$310,000", "Scale with tighter intent"],
                        ["Events", "$18,000", "104", "$220,000", "Selective"],
                        ["Partner", "$12,000", "92", "$205,000", "Efficiency test"],
                        ["Email", "$6,000", "84", "$98,000", "Nurture"],
                    ],
                    "widths": [1500, 1400, 1400, 1700, 3360],
                },
            },
        ],
    )
    _build_document(
        CALIBRATION_ROOT
        / "rwpk.evaluation_review_audit.quarterly_report_audit.0011"
        / "quarterly-report-audit.docx",
        title="Quarterly report audit",
        subtitle="Source mismatches and correction actions",
        kicker="Calibration reference",
        sections=[
            {
                "heading": "Critical findings",
                "table": {
                    "headers": ["Draft claim", "Source evidence", "Result"],
                    "rows": [
                        ["All regions improved Q2 orders", "South fell from 760 to 735", "Unsupported"],
                        ["Every region exceeded 94% on-time", "South 91%; West 93%", "Unsupported"],
                    ],
                    "widths": [3300, 3700, 2360],
                },
            },
            {
                "heading": "Required corrections",
                "bullets": [
                    "Replace the all-region growth claim with a region-specific summary.",
                    "State the 91%-95% on-time range.",
                    "Do not reuse the North playbook everywhere without causal evidence.",
                ],
            },
        ],
    )
    _build_document(
        CALIBRATION_ROOT
        / "rwpk.one_off_tool_execution.delivery_package.0012"
        / "handoff-note.docx",
        title="Quarterly delivery handoff",
        subtitle="Package contents, ownership, and remaining archive step",
        kicker="Calibration reference",
        sections=[
            {
                "heading": "Package status",
                "paragraphs": [
                    "The executive recap, delivery tracker, handoff note, and delivery manifest are present. The archive copy remains Ready for 2026-08-03 after final sign-off."
                ],
            },
            {
                "heading": "Ownership",
                "table": {
                    "headers": ["Item", "Owner role", "Status"],
                    "rows": [
                        ["Executive recap", "Program lead", "Complete"],
                        ["Delivery tracker", "Operations", "Complete"],
                        ["Handoff note", "Program lead", "Complete"],
                        ["Archive copy", "Operations", "Ready"],
                    ],
                    "widths": [3900, 2800, 2660],
                },
            },
        ],
    )


def _write_calibration_readme() -> None:
    _write_text(
        CALIBRATION_ROOT / "README.md",
        """
# Public calibration references

This directory contains 18 synthetic reference artifacts for the 12 v0.2 work
episodes. They demonstrate one valid way to satisfy each artifact contract and
are used to verify that the deterministic grader can recognize a conforming
package.

They are not byte-for-byte gold answers. Submissions may differ in structure,
wording, formulas, or visual design when the declared task properties, source
facts, filenames, editability requirements, and hard gates still pass.

## Human calibration status

Status: `not_run`

No practitioner score, model ranking, locale-parity conclusion, or leaderboard
claim is attached to these files. Before an episode can move from
`calibration_ready` to `reviewed`, record:

1. at least two blind workplace-practitioner reviews of materially different
   system outputs;
2. agreement and adjudication results for every human or model-judge item;
3. separate Korean and English calibration evidence;
4. any repaired or removed ambiguous rubric items;
5. reviewer roles, review date, protocol version, and immutable evidence path;
6. an independent run of the package without author assistance.

The scenario schema blocks `reviewed` and `frozen` status unless human
calibration is marked completed. Automated reference conformance alone cannot
promote an episode.
""",
    )


def prepare() -> None:
    if DATA_ROOT.exists():
        shutil.rmtree(DATA_ROOT)
    DATA_ROOT.mkdir(parents=True)
    _write_calibration_readme()
    _prepare_text_and_csv_assets()
    _prepare_documents()
    _build_vendor_pdf(
        ASSET_ROOT
        / "rwpk.summarization_synthesis.vendor_packet_brief.0001"
        / "vendor-packet.pdf"
    )
    _build_scanned_expense_pdf(
        ASSET_ROOT
        / "rwpk.extraction_parsing.scanned_expense_table.0002"
        / "scanned-expense-table.pdf"
    )
    print("prepared text, CSV, PDF, and DOCX fixtures")


def _media_type(path: Path) -> str:
    suffix = path.suffix.lower()
    return {
        ".pdf": "application/pdf",
        ".docx": DOCX_MIME,
        ".pptx": PPTX_MIME,
        ".xlsx": XLSX_MIME,
        ".csv": "text/csv",
        ".txt": "text/plain",
        ".json": "application/json",
    }[suffix]


def _artifact_family(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return "document"
    if suffix == ".pptx":
        return "presentation"
    if suffix in {".xlsx", ".csv"}:
        return "spreadsheet"
    if suffix == ".pdf":
        return "pdf_or_image"
    return "response_only"


def _asset(
    asset_id: str,
    path: Path,
    *,
    role: str,
    condition: str,
    trust: str = "trusted",
) -> dict[str, Any]:
    return {
        "asset_id": asset_id,
        "path": _relative(path),
        "media_type": _media_type(path),
        "artifact_family": _artifact_family(path),
        "attachment_role": role,
        "trust_level": trust,
        "input_condition": condition,
        "allowed_delivery_modes": ["native_file"],
        "locale_variant": "en-US",
        "locale_status": "shared_across_locales",
        "sha256": _sha256(path),
        "rights_basis": "original_project_authorship",
        "privacy": {
            "contains_personal_data": False,
            "classification": "public_synthetic",
        },
    }


def _check(check_id: str, check_type: str, **kwargs: Any) -> dict[str, Any]:
    return {"check_id": check_id, "type": check_type, **kwargs}


def _output(
    artifact_id: str,
    reference_path: Path,
    *,
    required_features: list[str],
    checks: list[dict[str, Any]],
) -> dict[str, Any]:
    return {
        "artifact_id": artifact_id,
        "artifact_family": _artifact_family(reference_path),
        "media_type": _media_type(reference_path),
        "required": True,
        "editable": reference_path.suffix.lower() in {".docx", ".pptx", ".xlsx"},
        "filename": reference_path.name,
        "required_features": required_features,
        "reference_path": _relative(reference_path),
        "reference_sha256": _sha256(reference_path),
        "reference_checks": checks,
    }


def _localized(ko: str, en: str) -> dict[str, str]:
    return {"ko": ko, "en": en}


def _realizations(
    scenario_id: str,
    assets: list[str],
    prompts: dict[str, str],
    naturalistic_features: list[str],
) -> list[dict[str, Any]]:
    result = []
    for locale, language_key in (("ko-KR", "ko"), ("en-US", "en")):
        for form in ("canonical", "naturalistic"):
            result.append(
                {
                    "prompt_id": f"{scenario_id}.{locale}.{form}",
                    "locale": locale,
                    "form": form,
                    "origin": (
                        "controlled_canonical"
                        if form == "canonical"
                        else "synthetic_naturalistic"
                    ),
                    "features": [] if form == "canonical" else naturalistic_features,
                    "messages": [
                        {
                            "role": "user",
                            "content": prompts[f"{language_key}_{form}"],
                            "attachment_refs": assets,
                        }
                    ],
                }
            )
    return result


def _verification(
    method: str,
    procedure: str,
    check: dict[str, Any] | None = None,
) -> dict[str, Any]:
    value: dict[str, Any] = {
        "method": method,
        "implementation_status": "implemented",
        "procedure": procedure,
    }
    if check is not None:
        value["check"] = check
    return value


def _evaluation(
    scenario_id: str,
    assets: list[dict[str, Any]],
    outputs: list[dict[str, Any]],
    *,
    response_mode: str = "answer_directly",
) -> dict[str, Any]:
    output_ids = [item["artifact_id"] for item in outputs]
    evidence_output = [
        {"source": "output", "ref": output_ids[0], "locator": "complete artifact bundle"}
    ]
    evidence_assets = [
        {"source": "asset", "ref": item["asset_id"], "locator": "whole supplied file"}
        for item in assets
    ]
    hard_gates = [
        {
            "gate_id": "hg.required_artifacts_present",
            "criterion": _localized(
                "필수 산출물이 모두 존재한다.",
                "Every required artifact is present.",
            ),
            "category": "artifact_presence",
            "artifact_refs": output_ids,
            "evidence_refs": evidence_output,
            "scorer": _verification(
                "deterministic",
                "Resolve every exact output filename in the submission directory.",
                _check("check.required_artifacts_present", "artifact_present"),
            ),
            "on_fail": "blocks_full_pass",
        },
        {
            "gate_id": "hg.media_types_match",
            "criterion": _localized(
                "각 산출물의 파일 형식이 계약과 일치한다.",
                "Every artifact media type matches the contract.",
            ),
            "category": "file_type",
            "artifact_refs": output_ids,
            "evidence_refs": evidence_output,
            "scorer": _verification(
                "deterministic",
                "Compare extension, file signature, and declared media type.",
                _check("check.media_types_match", "media_type"),
            ),
            "on_fail": "blocks_full_pass",
        },
        {
            "gate_id": "hg.artifacts_parse",
            "criterion": _localized(
                "모든 산출물을 참조 파서가 열 수 있다.",
                "Every artifact opens in the reference parser.",
            ),
            "category": "parseability",
            "artifact_refs": output_ids,
            "evidence_refs": evidence_output,
            "scorer": _verification(
                "deterministic",
                "Open ZIP-based Office files, JSON, and other formats with the reference inspector.",
                _check("check.artifacts_parse", "parseable"),
            ),
            "on_fail": "blocks_full_pass",
        },
        {
            "gate_id": "hg.required_structure_present",
            "criterion": _localized(
                "계약에 명시된 편집 가능한 구조가 존재한다.",
                "The required editable structure is present.",
            ),
            "category": "editability",
            "artifact_refs": output_ids,
            "evidence_refs": evidence_output,
            "scorer": _verification(
                "deterministic",
                "Inspect OOXML and JSON features declared by each output.",
                _check("check.required_structure_present", "required_features"),
            ),
            "on_fail": "blocks_full_pass",
        },
        {
            "gate_id": "hg.inputs_preserved",
            "criterion": _localized(
                "공개 입력 자산의 해시가 변경되지 않았다.",
                "The public input asset hashes remain unchanged.",
            ),
            "category": "input_preservation",
            "artifact_refs": [],
            "evidence_refs": evidence_assets,
            "scorer": _verification(
                "deterministic",
                "Recompute every declared input SHA-256 digest.",
                _check("check.inputs_preserved", "source_hashes_match"),
            ),
            "on_fail": "blocks_full_pass",
        },
    ]
    rubric_items: list[dict[str, Any]] = []
    for output in outputs:
        rubric_items.append(
            {
                "rubric_id": f"ar.{output['artifact_id']}_content",
                "criterion": _localized(
                    f"{output['filename']}에 작업별 핵심 사실과 구조가 포함된다.",
                    f"{output['filename']} contains the task-specific facts and structure.",
                ),
                "dimension": "source_fidelity",
                "requirement": "required",
                "scoring": {"type": "binary"},
                "scorer": _verification(
                    "rule_assisted",
                    "Run the output's declared reference checks against extracted structure and text.",
                    _check(
                        f"check.{output['artifact_id']}_reference",
                        "reference_checks",
                    ),
                ),
                "evidence_refs": [
                    {
                        "source": "output",
                        "ref": output["artifact_id"],
                        "locator": "declared reference checks",
                    }
                ],
                "adjudication_note": _localized(
                    "검사는 정답 파일의 바이트 일치가 아니라 요구 속성을 확인한다.",
                    "Checks verify required properties, not byte equality with the reference.",
                ),
            }
        )
    rubric_items.extend(
        [
            {
                "rubric_id": "ar.visual_quality",
                "criterion": _localized(
                    "렌더링된 산출물에 겹침, 잘림 또는 읽기 어려운 요소가 없다.",
                    "Rendered artifacts have no overlap, clipping, or unreadable elements.",
                ),
                "dimension": "visual_and_layout_quality",
                "requirement": "quality_improving",
                "scoring": {"type": "bounded", "minimum": 0, "maximum": 2, "step": 1},
                "scorer": _verification(
                    "human",
                    "Blindly inspect every rendered page, slide, or sheet at normal zoom.",
                ),
                "evidence_refs": evidence_output,
                "adjudication_note": _localized(
                    "공개 기준 산출물은 제작 QA를 통과했지만 모델 출력에 대한 인간 보정은 아직 필요하다.",
                    "Reference artifacts pass production QA; human calibration on model outputs remains required.",
                ),
            },
            {
                "rubric_id": "ar.scope_precision",
                "criterion": _localized(
                    "제출 폴더에 계약된 파일만 존재한다.",
                    "The submission folder contains only contracted files.",
                ),
                "dimension": "scope_precision",
                "requirement": "negative_check",
                "scoring": {"type": "binary"},
                "scorer": _verification(
                    "deterministic",
                    "Compare the submission file list with the exact output contract.",
                    _check(
                        "check.bundle_exact_files",
                        "bundle_exact_files",
                        values=[item["filename"] for item in outputs],
                    ),
                ),
                "evidence_refs": evidence_output,
                "adjudication_note": _localized(
                    "운영체제가 만든 숨김 파일은 무시한다.",
                    "Operating-system metadata files are ignored.",
                ),
            },
        ]
    )
    return {
        "response_mode": {
            "canonical": response_mode,
            "naturalistic": response_mode,
        },
        "hard_gates": hard_gates,
        "atomic_rubric_items": rubric_items,
        "human_calibration": {
            "required": True,
            "status": "not_run",
            "plan_path": "data/v0.2/public-calibration/README.md",
        },
    }


def _scenario(
    *,
    scenario_id: str,
    title: dict[str, str],
    intent: str,
    secondary_intents: list[str],
    workflow_job: str,
    domain: str,
    domain_tags: list[str],
    user_goal: dict[str, str],
    expected_artifact: str,
    module: str,
    input_conditions: list[str],
    naturalistic_features: list[str],
    assets: list[dict[str, Any]],
    outputs: list[dict[str, Any]],
    prompts: dict[str, str],
    response_mode: str = "answer_directly",
) -> dict[str, Any]:
    return {
        "schema": "realworld-prompt-kit.scenario/0.2.0",
        "scenario_id": scenario_id,
        "revision": 1,
        "status": "calibration_ready",
        "semantic_group_id": scenario_id.replace("rwpk.", "rwpg.", 1),
        "title": title,
        "task": {
            "primary_intent": intent,
            "secondary_intents": secondary_intents,
            "workflow_job": workflow_job,
            "primary_domain": domain,
            "domain_tags": domain_tags,
            "user_goal": user_goal,
            "expected_artifact": expected_artifact,
        },
        "coverage": {
            "module": module,
            "input_artifact_families": sorted(
                {item["artifact_family"] for item in assets}
            ),
            "output_artifact_families": sorted(
                {item["artifact_family"] for item in outputs}
            ),
            "input_conditions": input_conditions,
            "request_forms": ["canonical", "naturalistic"],
            "interaction_pattern": "one_shot",
            "authority": "local_reversible_write",
            "naturalistic_features": naturalistic_features,
        },
        "assets": assets,
        "realizations": _realizations(
            scenario_id,
            [item["asset_id"] for item in assets],
            prompts,
            naturalistic_features,
        ),
        "artifact_contract": {
            "outputs": outputs,
            "side_effect_scope": "local_reversible_write",
            "output_root": "outputs",
            "preserve_input_assets": True,
        },
        "evaluation": _evaluation(
            scenario_id,
            assets,
            outputs,
            response_mode=response_mode,
        ),
        "provenance": {
            "origin": "synthetic",
            "authors": ["RealWorld Prompt Kit"],
            "contains_personal_data": False,
            "rights_basis": "original_project_authorship",
            "license": "MIT",
            "review": {
                "review_type": "automated_structural",
                "reviewer": "Codex release QA",
                "review_date": BUILD_DATE,
                "evidence_path": "reports/v0.2/release-validation.json",
                "human_practitioner_review": False,
            },
        },
    }


def _asset_path(scenario_id: str, filename: str) -> Path:
    return ASSET_ROOT / scenario_id / filename


def _reference_path(scenario_id: str, filename: str) -> Path:
    return CALIBRATION_ROOT / scenario_id / filename


def _build_delivery_manifest() -> None:
    scenario_id = "rwpk.one_off_tool_execution.delivery_package.0012"
    files = [
        "delivery-recap-deck.pptx",
        "delivery-tracker.xlsx",
        "handoff-note.docx",
        "delivery-manifest.json",
    ]
    _write_json(
        _reference_path(scenario_id, "delivery-manifest.json"),
        {
            "schema": "realworld-prompt-kit.delivery-manifest/0.2.0",
            "package_id": "quarterly-delivery-handoff",
            "files": files,
            "authority": "local_reversible_write",
            "source_assets_preserved": True,
        },
    )


def _build_scenarios() -> list[dict[str, Any]]:
    scenarios: list[dict[str, Any]] = []

    sid = "rwpk.summarization_synthesis.vendor_packet_brief.0001"
    assets = [
        _asset(
            "vendor_packet",
            _asset_path(sid, "vendor-packet.pdf"),
            role="authoritative_source",
            condition="conflicting",
        )
    ]
    outputs = [
        _output(
            "vendor_brief",
            _reference_path(sid, "vendor-comparison-brief.docx"),
            required_features=["editable_text", "headings", "tables"],
            checks=[
                _check("check.vendor_text", "contains_text", values=["Cobalt", "$48,500", "$51,000", "conflict"]),
                _check("check.vendor_table", "minimum_feature", feature="table_count", minimum=1),
            ],
        )
    ]
    scenarios.append(
        _scenario(
            scenario_id=sid,
            title=_localized("공급업체 자료 비교 브리프", "Vendor packet comparison brief"),
            intent="summarization_synthesis",
            secondary_intents=["decision_recommendation"],
            workflow_job="synthesize",
            domain="supply_chain_logistics",
            domain_tags=["office_admin", "strategy_business_operations"],
            user_goal=_localized(
                "첨부 PDF의 상충하는 견적을 숨기지 않고 영어 의사결정 브리프로 정리한다.",
                "Turn the attached PDF into an English decision brief without hiding the conflicting quote.",
            ),
            expected_artifact="An editable English DOCX vendor comparison brief.",
            module="attached_file_grounding",
            input_conditions=["conflicting"],
            naturalistic_features=["missing_decisive_detail"],
            assets=assets,
            outputs=outputs,
            prompts={
                "ko_canonical": "첨부된 공급업체 PDF를 검토해 가격, 리드타임, 서비스 시간을 비교하는 영어 의사결정 브리프를 DOCX로 작성해 주세요. 상충하는 수치는 어느 쪽도 임의로 선택하지 말고 확인이 필요한 사항으로 표시해 주세요.",
                "ko_naturalistic": "이 업체 자료 영어 브리프로 정리해줘. 가격이 두 군데 다르게 적힌 것 같은데 그냥 하나 고르지 말고 뭐가 충돌하는지랑 다음 확인할 것까지 DOCX에 넣어줘.",
                "en_canonical": "Review the attached vendor PDF and create an editable English DOCX decision brief comparing price, lead time, and service hours. Do not silently choose between conflicting figures; identify the conflict and the required follow-up.",
                "en_naturalistic": "Can you turn this vendor packet into an English decision brief? I think one price shows up two ways, so flag the conflict instead of picking one and include the next check in the DOCX.",
            },
        )
    )

    sid = "rwpk.extraction_parsing.scanned_expense_table.0002"
    assets = [
        _asset(
            "expense_scan",
            _asset_path(sid, "scanned-expense-table.pdf"),
            role="authoritative_source",
            condition="noisy_or_scanned",
        )
    ]
    outputs = [
        _output(
            "expense_workbook",
            _reference_path(sid, "expense-table-extracted.xlsx"),
            required_features=["worksheets", "formulas", "editable_cells"],
            checks=[
                _check("check.expense_sheets", "required_sheet_names", values=["Extracted Expenses", "Checks"]),
                _check("check.expense_formulas", "minimum_feature", feature="formula_count", minimum=3),
                _check("check.expense_text", "contains_text", values=["R-1041", "R-1046", "Cobalt Freight"]),
            ],
        )
    ]
    scenarios.append(
        _scenario(
            scenario_id=sid,
            title=_localized("스캔 지출표 추출", "Scanned expense table extraction"),
            intent="extraction_parsing",
            secondary_intents=["quantitative_formal_analysis"],
            workflow_job="extract",
            domain="finance_accounting_tax",
            domain_tags=["office_admin", "data_analytics"],
            user_goal=_localized(
                "스캔 PDF 표를 검토 가능한 XLSX로 옮기고 합계 검사를 포함한다.",
                "Extract the scanned PDF table into a reviewable XLSX with checks.",
            ),
            expected_artifact="An editable XLSX with six extracted rows and formula checks.",
            module="attached_file_grounding",
            input_conditions=["noisy_or_scanned"],
            naturalistic_features=["ocr_copy_format_noise"],
            assets=assets,
            outputs=outputs,
            prompts={
                "ko_canonical": "첨부된 스캔 PDF의 지출표를 영어 열 이름을 유지한 편집 가능한 XLSX로 추출해 주세요. 여섯 행을 모두 보존하고 행 수, 금액 합계, 영수증 ID 중복 여부를 수식으로 확인하는 시트를 추가해 주세요.",
                "ko_naturalistic": "이 스캔 지출표 엑셀로 옮겨줘. 6줄 빠짐없이 넣고 합계랑 영수증 번호 중복 체크가 수식으로 보이게 해줘. 열 이름은 원문 영어 그대로면 돼.",
                "en_canonical": "Extract the attached scanned expense table into an editable XLSX. Preserve all six rows and add formula-driven checks for row count, amount total, and duplicate receipt IDs.",
                "en_naturalistic": "Please turn this scanned expense sheet into Excel. Keep all six lines and add formula checks for the total, row count, and duplicate receipt IDs.",
            },
        )
    )

    sid = "rwpk.communication_collaboration_negotiation.meeting_minutes.0003"
    assets = [
        _asset("meeting_notes", _asset_path(sid, "meeting-notes.txt"), role="authoritative_source", condition="clean"),
        _asset("meeting_agenda", _asset_path(sid, "meeting-agenda.docx"), role="supporting_source", condition="clean"),
    ]
    outputs = [
        _output(
            "meeting_minutes",
            _reference_path(sid, "program-sync-minutes.docx"),
            required_features=["editable_text", "headings", "tables"],
            checks=[
                _check("check.minutes_text", "contains_text", values=["2026-08-17", "Operations", "2026-08-05", "Unresolved questions"]),
                _check("check.minutes_table", "minimum_feature", feature="table_count", minimum=1),
            ],
        )
    ]
    scenarios.append(
        _scenario(
            scenario_id=sid,
            title=_localized("프로그램 회의록 작성", "Program sync minutes"),
            intent="communication_collaboration_negotiation",
            secondary_intents=["summarization_synthesis"],
            workflow_job="create",
            domain="communication_meetings",
            domain_tags=["project_product_management", "office_admin"],
            user_goal=_localized(
                "회의 노트와 안건으로 결정, 역할 소유자, 기한, 미해결 질문을 담은 영어 회의록을 만든다.",
                "Create English minutes with decisions, role owners, due dates, and unresolved questions.",
            ),
            expected_artifact="An editable English DOCX meeting-minutes document.",
            module="document_work",
            input_conditions=["clean"],
            naturalistic_features=["multi_intent_mixed_priority"],
            assets=assets,
            outputs=outputs,
            prompts={
                "ko_canonical": "첨부된 회의 노트와 안건을 사용해 영어 DOCX 회의록을 작성해 주세요. 결정, 역할 단위 소유자, 기한, 위험, 미해결 질문을 구분하고 제공되지 않은 개인 이름은 만들지 마세요.",
                "ko_naturalistic": "노트랑 아젠다로 영어 회의록 하나 만들어줘. 결정한 거, 누가 맡는지(사람 이름 말고 역할), 날짜, 아직 안 정한 거가 한눈에 보였으면 해.",
                "en_canonical": "Use the attached notes and agenda to create editable English DOCX minutes. Separate decisions, role-based owners, due dates, risks, and unresolved questions, and do not invent personal names.",
                "en_naturalistic": "Can you make English meeting minutes from the notes and agenda? I need the decisions, role owners, dates, and anything still unresolved to be easy to scan.",
            },
        )
    )

    sid = "rwpk.transformation_rewriting.onboarding_sop_revision.0004"
    assets = [
        _asset("sop_draft", _asset_path(sid, "onboarding-sop-draft.docx"), role="existing_artifact", condition="clean"),
        _asset("reviewer_feedback", _asset_path(sid, "reviewer-feedback.txt"), role="authoritative_source", condition="clean"),
    ]
    outputs = [
        _output(
            "revised_sop",
            _reference_path(sid, "onboarding-sop-revised.docx"),
            required_features=["editable_text", "headings", "tables"],
            checks=[
                _check("check.sop_text", "contains_text", values=["24 hours", "Hiring Manager", "Verify access", "Software procurement remains excluded"]),
                _check("check.sop_table", "minimum_feature", feature="table_count", minimum=1),
            ],
        )
    ]
    scenarios.append(
        _scenario(
            scenario_id=sid,
            title=_localized("온보딩 SOP 범위 수정", "Scoped onboarding SOP revision"),
            intent="transformation_rewriting",
            secondary_intents=["evaluation_review_audit"],
            workflow_job="revise",
            domain="hr_people_labor",
            domain_tags=["office_admin", "software_it"],
            user_goal=_localized(
                "기존 SOP를 피드백 범위 안에서만 수정하고 구조를 보존한다.",
                "Revise the existing SOP only within the reviewer-feedback scope.",
            ),
            expected_artifact="A minimally revised editable English DOCX SOP.",
            module="document_work",
            input_conditions=["clean"],
            naturalistic_features=["self_correction_scope_shift"],
            assets=assets,
            outputs=outputs,
            prompts={
                "ko_canonical": "첨부된 기존 SOP에 리뷰 피드백 네 항목만 반영해 영어 DOCX 개정본을 만들어 주세요. 24시간 목표, Hiring Manager 승인, IT Service Desk 검증 단계를 반영하고 소프트웨어 구매 제외 범위는 유지해 주세요.",
                "ko_naturalistic": "SOP 피드백대로만 고쳐줘. 아, 전체 재작성은 말고 기존 구성 살려서 24시간이랑 승인 역할, 마지막 IT 확인만 반영해. 구매 범위는 건드리지 마.",
                "en_canonical": "Apply only the four reviewer-feedback items to the existing SOP and produce an editable English DOCX. Preserve the structure, use the 24-hour target, assign approval to the Hiring Manager, add IT Service Desk verification, and retain the procurement exclusion.",
                "en_naturalistic": "Please update the SOP from the feedback, but don't rewrite everything. Keep the structure, change the target and approval role, add the IT check, and leave procurement out of scope.",
            },
        )
    )

    sid = "rwpk.summarization_synthesis.executive_update_deck.0005"
    assets = [
        _asset("portfolio_report", _asset_path(sid, "portfolio-report.docx"), role="authoritative_source", condition="clean"),
        _asset("workstream_metrics", _asset_path(sid, "workstream-metrics.csv"), role="authoritative_source", condition="clean"),
    ]
    outputs = [
        _output(
            "executive_deck",
            _reference_path(sid, "executive-update-deck.pptx"),
            required_features=["slides", "charts", "editable_text"],
            checks=[
                _check("check.executive_slides", "minimum_feature", feature="slide_count", minimum=3),
                _check("check.executive_chart", "minimum_feature", feature="chart_count", minimum=1),
                _check("check.executive_text", "contains_text", values=["3 / 4", "Enablement", "Support"]),
            ],
        )
    ]
    scenarios.append(
        _scenario(
            scenario_id=sid,
            title=_localized("임원 업데이트 덱", "Executive update deck"),
            intent="summarization_synthesis",
            secondary_intents=["decision_recommendation", "quantitative_formal_analysis"],
            workflow_job="create",
            domain="project_product_management",
            domain_tags=["strategy_business_operations", "data_analytics"],
            user_goal=_localized(
                "보고서와 지표 CSV에서 근거를 가져온 3장 영어 임원용 PPTX를 만든다.",
                "Create a three-slide English executive PPTX grounded in the report and metrics.",
            ),
            expected_artifact="An editable three-slide English executive update PPTX.",
            module="presentation_work",
            input_conditions=["clean"],
            naturalistic_features=["implicit_goal_or_output"],
            assets=assets,
            outputs=outputs,
            prompts={
                "ko_canonical": "첨부 보고서와 CSV를 근거로 3장짜리 영어 임원 업데이트 PPTX를 만들어 주세요. 진행 상황, 한 가지 핵심 리스크, 필요한 결정을 담고 수치를 원본과 일치시키며 편집 가능한 차트를 사용해 주세요.",
                "ko_naturalistic": "이 보고서랑 숫자로 임원 업데이트 덱 부탁해. 영어 3장이면 되고, 뭐가 잘 가는지랑 딱 필요한 결정이 보이게. 수치는 CSV랑 맞고 차트도 수정 가능해야 해.",
                "en_canonical": "Create a three-slide editable English executive update PPTX from the attached report and CSV. Show progress, the central risk, and required decisions; keep every figure source-grounded and use an editable chart.",
                "en_naturalistic": "Can you turn the report and numbers into a three-slide English exec update? Make the one real risk and the decisions obvious, keep the figures tied to the CSV, and use an editable chart.",
            },
        )
    )

    sid = "rwpk.transformation_rewriting.deck_feedback_revision.0006"
    assets = [
        _asset("existing_deck", _asset_path(sid, "existing-status-deck.pptx"), role="existing_artifact", condition="clean"),
        _asset("reviewer_comments", _asset_path(sid, "reviewer-comments.txt"), role="authoritative_source", condition="clean"),
    ]
    outputs = [
        _output(
            "revised_deck",
            _reference_path(sid, "revised-status-deck.pptx"),
            required_features=["slides", "charts", "editable_text"],
            checks=[
                _check("check.revised_slides", "minimum_feature", feature="slide_count", minimum=3),
                _check("check.revised_text", "contains_text", values=["76%", "handoff", "Q2 2026"]),
            ],
        )
    ]
    scenarios.append(
        _scenario(
            scenario_id=sid,
            title=_localized("피드백 기반 덱 개정", "Feedback-driven deck revision"),
            intent="transformation_rewriting",
            secondary_intents=["evaluation_review_audit"],
            workflow_job="revise",
            domain="strategy_business_operations",
            domain_tags=["project_product_management", "communication_meetings"],
            user_goal=_localized(
                "기존 덱에 지정 피드백을 적용하고 관련 없는 범위와 원본 수치를 보존한다.",
                "Apply specified feedback while preserving unrelated scope and source figures.",
            ),
            expected_artifact="A scoped three-slide revision of the English PPTX.",
            module="presentation_work",
            input_conditions=["clean"],
            naturalistic_features=["mid_task_change"],
            assets=assets,
            outputs=outputs,
            prompts={
                "ko_canonical": "기존 PPTX에 첨부된 리뷰 코멘트를 반영해 영어 개정본을 만들어 주세요. 검증된 준비도와 자체 보고를 구분하고 handoff 책임을 핵심 리스크로 만들되, Q2 2026과 원본 수치는 유지하고 3장 범위를 넘기지 마세요.",
                "ko_naturalistic": "이 덱 코멘트대로 손봐줘. 준비됐다고 한 거랑 실제 확인된 거는 나누고 handoff 담당 문제를 중심에 둬. 아, 숫자나 Q2 표시는 바꾸지 말고 3장 그대로.",
                "en_canonical": "Revise the existing English PPTX using the reviewer comments. Distinguish verified readiness from self-reporting, make handoff ownership the central risk, preserve Q2 2026 and source figures, and keep the deck to three slides.",
                "en_naturalistic": "Please update this deck from the comments. Separate what's verified from what's just reported and center the handoff-owner risk. Keep the Q2 label, figures, and three-slide scope.",
            },
        )
    )

    sid = "rwpk.quantitative_formal_analysis.project_status_workbook.0007"
    assets = [
        _asset("project_status", _asset_path(sid, "project-status.csv"), role="authoritative_source", condition="clean")
    ]
    outputs = [
        _output(
            "status_workbook",
            _reference_path(sid, "project-status-summary.xlsx"),
            required_features=["worksheets", "formulas", "charts", "editable_cells"],
            checks=[
                _check("check.status_sheets", "required_sheet_names", values=["Source Data", "Status Summary", "Assumptions"]),
                _check("check.status_formulas", "minimum_feature", feature="formula_count", minimum=6),
                _check("check.status_text", "contains_text", values=["On Track", "At Risk", "Blocked"]),
            ],
        )
    ]
    scenarios.append(
        _scenario(
            scenario_id=sid,
            title=_localized("프로젝트 상태 워크북", "Project status workbook"),
            intent="quantitative_formal_analysis",
            secondary_intents=["classification_organization"],
            workflow_job="create",
            domain="project_product_management",
            domain_tags=["strategy_business_operations", "data_analytics"],
            user_goal=_localized(
                "프로젝트 CSV를 보존하고 상태별 건수와 예산을 수식으로 요약한다.",
                "Preserve project CSV rows and summarize counts and budgets with formulas.",
            ),
            expected_artifact="An editable XLSX with source, summary, assumptions, formulas, and chart.",
            module="spreadsheet_work",
            input_conditions=["clean"],
            naturalistic_features=["colloquial_slang_dialect"],
            assets=assets,
            outputs=outputs,
            prompts={
                "ko_canonical": "첨부 CSV의 모든 원본 행을 별도 시트에 보존하고 상태별 프로젝트 수, 예산, 사용액을 수식으로 계산하는 편집 가능한 XLSX를 만들어 주세요. 가정 시트와 상태별 차트를 포함해 주세요.",
                "ko_naturalistic": "이 CSV로 상태판 엑셀 하나 만들어줘. 원본은 따로 두고 상태별 몇 건인지랑 예산/사용액이 수식으로 보였으면 해. 가정이랑 차트도 넣어줘.",
                "en_canonical": "Create an editable XLSX from the attached CSV. Preserve every source row, calculate project counts, budget, and budget used by status with formulas, and include an assumptions sheet and status chart.",
                "en_naturalistic": "Can you make a status workbook from this CSV? Keep the raw rows, show counts and budget totals by status with formulas, and add assumptions plus a chart.",
            },
        )
    )

    sid = "rwpk.diagnosis_root_cause.budget_workbook_repair.0008"
    assets = [
        _asset("broken_budget", _asset_path(sid, "budget-with-range-error.xlsx"), role="existing_artifact", condition="incomplete")
    ]
    outputs = [
        _output(
            "repaired_budget",
            _reference_path(sid, "budget-repaired.xlsx"),
            required_features=["worksheets", "formulas", "editable_cells"],
            checks=[
                _check("check.budget_sheets", "required_sheet_names", values=["Budget", "Checks"]),
                _check("check.budget_formulas", "minimum_feature", feature="formula_count", minimum=8),
                _check("check.budget_text", "contains_text", values=["Technology", "Variance reconciliation"]),
            ],
        )
    ]
    scenarios.append(
        _scenario(
            scenario_id=sid,
            title=_localized("예산 워크북 수식 복구", "Budget workbook formula repair"),
            intent="diagnosis_root_cause",
            secondary_intents=["quantitative_formal_analysis", "evaluation_review_audit"],
            workflow_job="validate",
            domain="finance_accounting_tax",
            domain_tags=["office_admin", "data_analytics"],
            user_goal=_localized(
                "누락된 합계 범위를 고치고 검증 시트를 추가하되 원본 입력을 보존한다.",
                "Repair the omitted total range and add checks while preserving inputs.",
            ),
            expected_artifact="An editable repaired XLSX with reconciled totals and a checks sheet.",
            module="spreadsheet_work",
            input_conditions=["incomplete"],
            naturalistic_features=["frustration_urgency_emotion"],
            assets=assets,
            outputs=outputs,
            prompts={
                "ko_canonical": "첨부 예산 워크북의 합계와 분산 수식을 점검해 누락된 범위를 수정하고 영어 XLSX 개정본을 만들어 주세요. 입력 값과 기존 형식은 유지하고, 합계 및 분산 조정을 검증하는 Checks 시트를 추가해 주세요.",
                "ko_naturalistic": "이 예산표 total이 계속 안 맞아. 값은 건드리지 말고 빠진 범위 찾아서 수식 고쳐줘. Checks 시트에서 합계랑 variance가 맞는지도 보이게 해줘.",
                "en_canonical": "Audit the attached budget workbook, repair the omitted range in the totals, and create an editable English XLSX revision. Preserve inputs and formatting and add a Checks sheet that reconciles totals and variance.",
                "en_naturalistic": "This budget total won't tie. Please keep the inputs, fix the missing formula range, and add a Checks sheet showing that the totals and variance reconcile.",
            },
        )
    )

    sid = "rwpk.operations_monitoring_improvement.monthly_ops_package.0009"
    assets = [
        _asset("monthly_data", _asset_path(sid, "monthly-ops.csv"), role="authoritative_source", condition="clean"),
        _asset("management_brief", _asset_path(sid, "management-brief.txt"), role="authoritative_source", condition="clean"),
    ]
    outputs = [
        _output(
            "ops_workbook",
            _reference_path(sid, "monthly-ops-workbook.xlsx"),
            required_features=["worksheets", "formulas", "charts", "editable_cells"],
            checks=[
                _check("check.ops_sheets", "required_sheet_names", values=["Monthly Data", "Management Summary"]),
                _check("check.ops_workbook_text", "contains_text", values=["Jun", "1525", "Backlog"]),
            ],
        ),
        _output(
            "ops_memo",
            _reference_path(sid, "monthly-ops-memo.docx"),
            required_features=["editable_text", "headings"],
            checks=[
                _check("check.ops_memo_text", "contains_text", values=["1,525", "95%", "54", "93%"]),
            ],
        ),
        _output(
            "ops_deck",
            _reference_path(sid, "monthly-ops-deck.pptx"),
            required_features=["slides", "charts", "editable_text"],
            checks=[
                _check("check.ops_deck_slides", "minimum_feature", feature="slide_count", minimum=3),
                _check("check.ops_deck_text", "contains_text", values=["95%", "1,525", "54"]),
            ],
        ),
    ]
    scenarios.append(
        _scenario(
            scenario_id=sid,
            title=_localized("월간 운영 관리 패키지", "Monthly operations management package"),
            intent="operations_monitoring_improvement",
            secondary_intents=["summarization_synthesis", "quantitative_formal_analysis"],
            workflow_job="package",
            domain="strategy_business_operations",
            domain_tags=["data_analytics", "communication_meetings"],
            user_goal=_localized(
                "한 원본에서 일관된 워크북, 메모, 덱을 만든다.",
                "Create a consistent workbook, memo, and deck from one source.",
            ),
            expected_artifact="A coordinated English XLSX, DOCX, and PPTX package.",
            module="cross_artifact_workflow",
            input_conditions=["clean"],
            naturalistic_features=["multi_intent_mixed_priority"],
            assets=assets,
            outputs=outputs,
            prompts={
                "ko_canonical": "첨부 CSV와 관리 브리프를 사용해 영어 관리 패키지를 만들어 주세요. 원본과 수식 요약이 있는 XLSX, 1~2쪽 DOCX 메모, 3장 PPTX를 만들고 June의 주문, 정시율, backlog 수치를 세 파일에서 일치시켜 주세요.",
                "ko_naturalistic": "이 월간 데이터로 관리 패키지 부탁해. 엑셀 원본/요약, 짧은 영어 메모, 3장 덱까지. 특히 June 주문이랑 on-time, backlog 숫자는 셋 다 똑같아야 해.",
                "en_canonical": "Use the attached CSV and management brief to create an English management package: an XLSX with source data and formula summary, a one-to-two-page DOCX memo, and a three-slide PPTX. Keep June orders, on-time rate, and backlog consistent across all files.",
                "en_naturalistic": "Can you package this monthly data into a workbook, short English memo, and three-slide deck? Keep the June orders, on-time rate, and backlog identical across all three.",
            },
        )
    )

    sid = "rwpk.evaluation_review_audit.campaign_readout_package.0010"
    assets = [
        _asset("campaign_results", _asset_path(sid, "campaign-results.xlsx"), role="authoritative_source", condition="clean"),
        _asset("brand_guide", _asset_path(sid, "brand-guide.docx"), role="style_reference", condition="clean"),
    ]
    outputs = [
        _output(
            "campaign_brief",
            _reference_path(sid, "campaign-readout-brief.docx"),
            required_features=["editable_text", "headings", "tables"],
            checks=[
                _check("check.campaign_brief_text", "contains_text", values=["Partner", "$130", "Search", "$310,000"]),
            ],
        ),
        _output(
            "campaign_deck",
            _reference_path(sid, "campaign-recap-deck.pptx"),
            required_features=["slides", "charts", "editable_text"],
            checks=[
                _check("check.campaign_deck_slides", "minimum_feature", feature="slide_count", minimum=3),
                _check("check.campaign_deck_text", "contains_text", values=["Partner", "$130", "Search"]),
            ],
        ),
    ]
    scenarios.append(
        _scenario(
            scenario_id=sid,
            title=_localized("캠페인 리드아웃 재구성", "Campaign readout repackaging"),
            intent="evaluation_review_audit",
            secondary_intents=["decision_recommendation", "quantitative_formal_analysis"],
            workflow_job="repurpose",
            domain="sales_marketing",
            domain_tags=["data_analytics", "strategy_business_operations"],
            user_goal=_localized(
                "캠페인 워크북을 브랜드 지침에 맞는 브리프와 덱으로 재구성한다.",
                "Repurpose the campaign workbook into a branded brief and deck.",
            ),
            expected_artifact="An English DOCX brief and three-slide PPTX with consistent findings.",
            module="cross_artifact_workflow",
            input_conditions=["clean"],
            naturalistic_features=["code_switching_jargon"],
            assets=assets,
            outputs=outputs,
            prompts={
                "ko_canonical": "캠페인 결과 XLSX와 브랜드 가이드를 바탕으로 영어 DOCX 리드아웃과 3장 PPTX를 만들어 주세요. 규모와 효율을 구분하고 Partner와 Search의 결론을 두 파일에서 일치시키며 제공된 색상과 직접적인 문체를 사용해 주세요.",
                "ko_naturalistic": "이 campaign 결과로 영어 brief랑 3장 deck 만들어줘. scale이랑 efficiency는 구분하고 Partner/Search 얘기가 두 파일에서 안 엇갈리게. brand guide 톤이랑 컬러도 맞춰줘.",
                "en_canonical": "Use the campaign-results XLSX and brand guide to create an English DOCX readout and three-slide PPTX. Distinguish scale from efficiency, keep the Partner and Search conclusions consistent, and follow the supplied tone and colors.",
                "en_naturalistic": "Please turn the campaign workbook into an English brief and three-slide deck. Separate scale from efficiency, keep the Partner/Search story consistent, and follow the brand guide.",
            },
        )
    )

    sid = "rwpk.evaluation_review_audit.quarterly_report_audit.0011"
    assets = [
        _asset("draft_report", _asset_path(sid, "quarterly-report-draft.docx"), role="existing_artifact", condition="conflicting"),
        _asset("source_data", _asset_path(sid, "quarterly-source.csv"), role="authoritative_source", condition="clean"),
    ]
    outputs = [
        _output(
            "audit_report",
            _reference_path(sid, "quarterly-report-audit.docx"),
            required_features=["editable_text", "headings", "tables"],
            checks=[
                _check("check.audit_text", "contains_text", values=["South", "760", "735", "91%", "Unsupported"]),
                _check("check.audit_table", "minimum_feature", feature="table_count", minimum=1),
            ],
        )
    ]
    scenarios.append(
        _scenario(
            scenario_id=sid,
            title=_localized("분기 보고서 근거 감사", "Quarterly report evidence audit"),
            intent="evaluation_review_audit",
            secondary_intents=["quantitative_formal_analysis"],
            workflow_job="inspect",
            domain="strategy_business_operations",
            domain_tags=["data_analytics", "office_admin"],
            user_goal=_localized(
                "초안 보고서의 근거 오류를 CSV와 대조해 영어 감사 보고서로 기록한다.",
                "Audit draft claims against the CSV and document the findings.",
            ),
            expected_artifact="An editable English DOCX audit with claim-level evidence.",
            module="artifact_quality_control_delivery",
            input_conditions=["conflicting"],
            naturalistic_features=["contradictory_constraints"],
            assets=assets,
            outputs=outputs,
            prompts={
                "ko_canonical": "분기 보고서 초안을 CSV와 대조해 영어 DOCX 감사 보고서를 작성해 주세요. 잘못된 전체 지역 성장 및 정시율 주장을 근거 행과 함께 표시하고, 원본 초안 자체는 수정하지 말며 필요한 교정만 제안해 주세요.",
                "ko_naturalistic": "이 보고서 숫자 맞는지 CSV랑 대조해줘. 초안 파일은 손대지 말고 영어 audit 문서로 어떤 문장이 틀렸는지, 실제 숫자랑 고칠 방향만 적어줘.",
                "en_canonical": "Audit the draft quarterly report against the attached CSV and create an editable English DOCX audit. Identify the unsupported all-region growth and on-time-rate claims with source evidence; preserve the draft and recommend only necessary corrections.",
                "en_naturalistic": "Can you check this report against the CSV? Don't edit the draft itself; give me an English audit showing which claims are wrong, the actual figures, and the needed corrections.",
            },
        )
    )

    sid = "rwpk.one_off_tool_execution.delivery_package.0012"
    assets = [
        _asset("delivery_brief", _asset_path(sid, "delivery-brief.docx"), role="authoritative_source", condition="clean"),
        _asset("delivery_data", _asset_path(sid, "delivery-data.csv"), role="authoritative_source", condition="clean"),
    ]
    outputs = [
        _output(
            "delivery_deck",
            _reference_path(sid, "delivery-recap-deck.pptx"),
            required_features=["slides", "charts", "editable_text"],
            checks=[
                _check("check.delivery_deck_text", "contains_text", values=["4 / 4", "August 3", "handoff"]),
            ],
        ),
        _output(
            "delivery_tracker",
            _reference_path(sid, "delivery-tracker.xlsx"),
            required_features=["worksheets", "formulas", "editable_cells"],
            checks=[
                _check("check.delivery_tracker_sheet", "required_sheet_names", values=["Delivery Tracker"]),
                _check("check.delivery_tracker_text", "contains_text", values=["Archive copy", "Ready"]),
            ],
        ),
        _output(
            "handoff_note",
            _reference_path(sid, "handoff-note.docx"),
            required_features=["editable_text", "headings", "tables"],
            checks=[
                _check("check.handoff_text", "contains_text", values=["2026-08-03", "Archive copy", "Ready"]),
            ],
        ),
        _output(
            "delivery_manifest",
            _reference_path(sid, "delivery-manifest.json"),
            required_features=["json_object", "file_manifest"],
            checks=[
                _check("check.manifest_keys", "json_keys", values=["schema", "package_id", "files", "authority", "source_assets_preserved"]),
                _check("check.manifest_files", "contains_text", values=["delivery-recap-deck.pptx", "delivery-tracker.xlsx", "handoff-note.docx"]),
            ],
        ),
    ]
    scenarios.append(
        _scenario(
            scenario_id=sid,
            title=_localized("분기 인계 패키지", "Quarterly handoff package"),
            intent="one_off_tool_execution",
            secondary_intents=["communication_collaboration_negotiation", "classification_organization"],
            workflow_job="package",
            domain="office_admin",
            domain_tags=["project_product_management", "strategy_business_operations"],
            user_goal=_localized(
                "정확한 파일명과 로컬 쓰기 범위로 영어 인계 패키지를 만든다.",
                "Create an English handoff package with exact filenames and local-only writes.",
            ),
            expected_artifact="A four-file local package: PPTX, XLSX, DOCX, and JSON manifest.",
            module="artifact_quality_control_delivery",
            input_conditions=["clean"],
            naturalistic_features=["implicit_permission_or_authority"],
            assets=assets,
            outputs=outputs,
            prompts={
                "ko_canonical": "첨부 brief와 CSV로 영어 분기 인계 패키지를 만들어 주세요. 지정된 이름의 3장 PPTX, 편집 가능한 XLSX tracker, DOCX handoff note, JSON manifest를 로컬 출력 폴더에만 작성하고 입력 파일은 변경하지 마세요.",
                "ko_naturalistic": "이 자료로 handoff 묶음 마무리해줘. 영어 덱/트래커/노트랑 manifest까지 brief에 적힌 파일명 그대로. 로컬 폴더에만 만들고 원본은 건드리지 마.",
                "en_canonical": "Use the attached brief and CSV to create the English quarterly handoff package: a three-slide PPTX, editable XLSX tracker, DOCX handoff note, and JSON manifest with the exact required filenames. Write only to the local output folder and preserve the inputs.",
                "en_naturalistic": "Please finish the handoff bundle from these files: English deck, tracker, note, and manifest with the names in the brief. Keep it local and don't touch the source files.",
            },
        )
    )
    return scenarios


def _build_manifest_and_catalog(scenarios: list[dict[str, Any]]) -> None:
    catalog = {
        "schema": "realworld-prompt-kit.catalog/0.2.0",
        "version": "0.2.0",
        "track": "artifact_core",
        "description": "Twelve portable, bilingual artifact-grounded office-work calibration episodes.",
        "workflow_jobs": WORKFLOW_JOBS,
        "modules": [
            {
                "module_id": module,
                "episode_count": sum(
                    scenario["coverage"]["module"] == module for scenario in scenarios
                ),
            }
            for module in MODULES
        ],
        "artifact_families": [
            "response_only",
            "pdf_or_image",
            "document",
            "presentation",
            "spreadsheet",
            "mixed_bundle",
        ],
        "input_conditions": [
            "clean",
            "noisy_or_scanned",
            "incomplete",
            "conflicting",
            "untrusted_embedded_content",
        ],
        "interaction_patterns": [
            "one_shot",
            "clarification_required",
            "feedback_revision",
            "scope_change",
            "infeasible_or_hold",
        ],
        "episodes": [
            {
                "scenario_id": scenario["scenario_id"],
                "path": _relative(
                    SCENARIO_ROOT / f"{scenario['scenario_id']}.json"
                ),
                "module": scenario["coverage"]["module"],
                "workflow_job": scenario["task"]["workflow_job"],
                "primary_intent": scenario["task"]["primary_intent"],
                "output_artifact_families": scenario["coverage"][
                    "output_artifact_families"
                ],
            }
            for scenario in scenarios
        ],
    }
    manifest = {
        "schema": "realworld-prompt-kit.manifest/0.2.0",
        "pack_id": "realworld-prompt-kit-v0.2-artifact-core",
        "version": "0.2.0",
        "status": "calibration_release",
        "description": "The v0.2 Artifact Core calibration release: 12 complete work episodes with native assets, editable references, atomic rubrics, and deterministic grading.",
        "catalog_path": "data/v0.2/catalog.json",
        "scenario_glob": "data/v0.2/scenarios/*.json",
        "expected_scenarios": 12,
        "expected_realizations": 48,
        "realizations_per_scenario": 4,
        "required_locales": ["ko-KR", "en-US"],
        "required_forms": ["canonical", "naturalistic"],
        "required_modules": {module: 2 for module in MODULES},
        "required_workflow_jobs": WORKFLOW_JOBS,
        "required_output_families": ["document", "presentation", "spreadsheet"],
        "reference_validation": {
            "required": True,
            "all_hard_gates_must_pass": True,
            "report_path": "reports/v0.2/release-validation.json",
        },
        "source_policy": {
            "allowed_origin": "synthetic",
            "require_personal_data_false": True,
            "rights_basis": "original_project_authorship",
            "license": "MIT",
        },
        "claim_boundaries": {
            "leaderboard_valid": False,
            "human_calibration_complete": False,
            "external_replication_complete": False,
            "supported_claim": "portable artifact construction and deterministic contract conformance on the public calibration episodes",
            "unsupported_claims": [
                "global office-worker capability ranking",
                "population task frequency",
                "GUI-operation proficiency",
                "stateful external action safety",
            ],
        },
        "release_evidence": {
            "build_date": BUILD_DATE,
            "scenario_schema": "schemas/scenario-0.2.schema.json",
            "artifact_validator": "tools/artifacts/validate.py",
            "artifact_grader": "tools/artifacts/grade_artifacts.py",
            "reference_outputs": "data/v0.2/public-calibration/",
            "human_calibration_plan": "data/v0.2/public-calibration/README.md",
        },
    }
    _write_json(DATA_ROOT / "catalog.json", catalog)
    _write_json(DATA_ROOT / "manifest.json", manifest)


def finalize() -> None:
    required_office_files = [
        _reference_path(
            "rwpk.summarization_synthesis.executive_update_deck.0005",
            "executive-update-deck.pptx",
        ),
        _reference_path(
            "rwpk.quantitative_formal_analysis.project_status_workbook.0007",
            "project-status-summary.xlsx",
        ),
        _asset_path(
            "rwpk.diagnosis_root_cause.budget_workbook_repair.0008",
            "budget-with-range-error.xlsx",
        ),
    ]
    missing = [str(path) for path in required_office_files if not path.is_file()]
    if missing:
        raise SystemExit(
            "office builder outputs are missing; run build_office_assets.mjs first:\n"
            + "\n".join(missing)
        )
    _build_delivery_manifest()
    for path in sorted(DATA_ROOT.rglob("*")):
        if path.is_file() and path.suffix.lower() in {".docx", ".pptx", ".xlsx"}:
            _normalize_ooxml(path)
    scenarios = _build_scenarios()
    SCENARIO_ROOT.mkdir(parents=True, exist_ok=True)
    for scenario in scenarios:
        _write_json(
            SCENARIO_ROOT / f"{scenario['scenario_id']}.json",
            scenario,
        )
    _build_manifest_and_catalog(scenarios)
    print(
        f"finalized v0.2 release metadata: {len(scenarios)} scenarios, "
        f"{sum(len(item['realizations']) for item in scenarios)} realizations"
    )


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--phase", choices=["prepare", "finalize"], required=True)
    args = parser.parse_args()
    if args.phase == "prepare":
        prepare()
    else:
        finalize()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
